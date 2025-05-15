# PDF에서 텍스트를 추출하는 코드 예제입니다:

import pdfplumber

def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    return text

# SQLAlchemy를 사용해 데이터를 저장하고 검색합니다:
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import datetime

# 데이터베이스 설정
Base = declarative_base()

class Proposal(Base):
    __tablename__ = 'proposals'
    id = Column(Integer, primary_key=True, autoincrement=True)
    title = Column(String, nullable=False)
    content = Column(Text, nullable=False)
    tags = Column(String, nullable=True)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

# 데이터베이스 연결
engine = create_engine('sqlite:///proposals.db')
Base.metadata.create_all(engine)
Session = sessionmaker(bind=engine)
session = Session()

def search_similar_proposals(keyword):
    return session.query(Proposal).filter(Proposal.tags.contains(keyword)).all()

# SQLAlchemy를 사용해 데이터를 저장하고 검색합니다:
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import datetime

# 데이터베이스 설정
Base = declarative_base()

class Proposal(Base):
    __tablename__ = 'proposals'
    id = Column(Integer, primary_key=True, autoincrement=True)
    title = Column(String, nullable=False)
    content = Column(Text, nullable=False)
    tags = Column(String, nullable=True)
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

# 데이터베이스 연결
engine = create_engine('sqlite:///proposals.db')
Base.metadata.create_all(engine)
Session = sessionmaker(bind=engine)
session = Session()

def search_similar_proposals(keyword):
    return session.query(Proposal).filter(Proposal.tags.contains(keyword)).all()

# 템플릿 기반 제안서 생성
from docx import Document

def create_proposal(rfp_data, similar_proposals):
    doc = Document()
    doc.add_heading('감리 사업 제안서', level=1)
    doc.add_paragraph(f"프로젝트명: {rfp_data['project_name']}")
    doc.add_paragraph(f"감리 범위: {rfp_data['audit_scope']}")
    doc.add_paragraph(f"시작일: {rfp_data['start_date']}")
    doc.add_paragraph(f"종료일: {rfp_data['end_date']}")
    
    # 유사 제안서 내용 참고
    doc.add_heading('참고 내용', level=2)
    for proposal in similar_proposals:
        doc.add_paragraph(f"- {proposal.title}: {proposal.content[:100]}...")

    output_file = "generated_proposal.docx"
    doc.save(output_file)
    return output_file

# 합 프로그램
def main(rfp_file):
    # RFP 파일 분석
    rfp_text = extract_text_from_pdf(rfp_file)
    rfp_data = {
        "project_name": "Sample Project",
        "audit_scope": "Compliance and Security Review",
        "start_date": "2024-01-01",
        "end_date": "2024-03-31"
    }

    # 키워드로 유사 제안서 검색
    keyword = "Security"
    similar_proposals = search_similar_proposals(keyword)

    # 제안서 생성
    output_file = create_proposal(rfp_data, similar_proposals)
    print(f"제안서가 생성되었습니다: {output_file}")

if __name__ == "__main__":
    main("sample_rfp.pdf")
